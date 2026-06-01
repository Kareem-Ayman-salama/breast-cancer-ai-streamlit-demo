from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"E:\Freelance work\nour")
OUT = ROOT / "outputs"
DELIVERABLES = ROOT / "client_deliverables"
DOCX_PATH = DELIVERABLES / "breast_cancer_ai_client_report.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(23, 32, 42)
MUTED = RGBColor(90, 102, 112)
HEADER_FILL = "F2F4F7"
CALLOUT_FILL = "FFF8E1"
BORDER = "D9E2EC"


def load_json(path: Path) -> dict:
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def read_csv(path: Path, **kwargs) -> pd.DataFrame:
    if not path.exists():
        return pd.DataFrame()
    return pd.read_csv(path, **kwargs)


def fmt(value, digits: int = 3) -> str:
    try:
        return f"{float(value):.{digits}f}"
    except Exception:
        return str(value) if value is not None else "NA"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa=9360, indent_dxa=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    run.italic = True


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.1) -> None:
    if not path.exists():
        p = doc.add_paragraph()
        run = p.add_run(f"[Missing figure: {caption}]")
        run.font.color.rgb = RGBColor(155, 28, 28)
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, caption)


def add_table_from_df(
    doc: Document,
    df: pd.DataFrame,
    widths: list[int] | None = None,
    font_size: float = 8.5,
    max_rows: int | None = None,
) -> None:
    if df.empty:
        doc.add_paragraph("No data available.")
        return
    table_df = df.copy()
    if max_rows is not None:
        table_df = table_df.head(max_rows)
    table_df = table_df.fillna("NA")

    table = doc.add_table(rows=1, cols=len(table_df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table)
    set_table_borders(table)

    if widths is None:
        widths = [int(9360 / len(table_df.columns))] * len(table_df.columns)

    header_cells = table.rows[0].cells
    for idx, col in enumerate(table_df.columns):
        cell = header_cells[idx]
        set_cell_shading(cell, HEADER_FILL)
        set_cell_margins(cell)
        set_cell_width(cell, widths[idx])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(col))
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(font_size)
        run.font.color.rgb = INK

    for _, row in table_df.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(table_df.columns):
            cell = cells[idx]
            set_cell_margins(cell)
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            text = str(row[col])
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if len(text) > 14 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.name = "Calibri"
            run.font.size = Pt(font_size)
            run.font.color.rgb = INK

    doc.add_paragraph()


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = INK


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table)
    set_table_borders(table, color="F59E0B", size="6")
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    r1 = p.add_run(title + ": ")
    r1.bold = True
    r1.font.name = "Calibri"
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = RGBColor(122, 90, 0)
    r2 = p.add_run(body)
    r2.font.name = "Calibri"
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = INK
    doc.add_paragraph()


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    bullet = doc.styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet.font.size = Pt(10.5)
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing = 1.167


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.text = "Breast Cancer AI Project"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if p.runs:
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = MUTED

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Client report")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def build_overview_rows() -> pd.DataFrame:
    tcga = load_json(OUT / "tcga_brca_cells" / "cancer_prediction" / "cancer_holdout_metrics.json")
    bcsc = load_json(OUT / "bcsc_future_risk_prediction" / "bcsc_future_risk_metrics.json")
    subtype = load_json(OUT / "tumor_subtype_prediction" / "tumor_subtype_metrics.json")
    survival = load_json(OUT / "metabric_optimized_prognosis" / "survival_best_metrics.json")
    recurrence = load_json(OUT / "metabric_optimized_prognosis" / "recurrence_best_metrics.json")
    corr = load_json(OUT / "gene_correlation_pattern_analysis" / "gene_correlation_pattern_summary.json")
    upload_metrics = load_json(OUT / "streamlit_upload_gene_models" / "upload_gene_models_metrics.json")
    external = read_csv(OUT / "geo_external_validation" / "metabric_to_geo_external_validation_metrics.csv")

    external_auc = "NA"
    if not external.empty and "roc_auc" in external.columns:
        external_auc = f"{external['roc_auc'].min():.3f} to {external['roc_auc'].max():.3f}"

    rows = [
        ["Cancer molecular signature", "TCGA-BRCA", "Cancer vs Normal tissue", tcga.get("samples", "NA"), f"AUC {fmt(tcga.get('roc_auc'))}, Acc {fmt(tcga.get('accuracy'))}"],
        ["Future incidence risk", "BCSC", "1-year future diagnosis risk", int(bcsc.get("weighted_training_records", 0) + bcsc.get("weighted_validation_records", 0)), f"AUC {fmt(bcsc.get('validation', {}).get('roc_auc'))}"],
        ["Tumor subtype", "METABRIC", "Real CLAUDIN subtype", subtype.get("samples_predicted", "NA"), f"Balanced Acc {fmt(subtype.get('holdout_balanced_accuracy_from_previous_training'))}"],
        ["Survival prognosis", "METABRIC", "Overall survival event", survival.get("samples", "NA"), f"Holdout AUC {fmt(survival.get('holdout_roc_auc'))}"],
        ["Recurrence prognosis", "METABRIC", "Recurrence event", recurrence.get("samples", "NA"), f"Holdout AUC {fmt(recurrence.get('holdout_roc_auc'))}"],
        ["External validation", "GEO", "Relapse, metastasis, death, RFS, OS", int(external["samples"].sum()) if not external.empty else "NA", f"AUC range {external_auc}"],
        [
            "Gene-expression upload",
            "Streamlit demo",
            "User-uploaded gene CSV prediction",
            upload_metrics.get("survival", {}).get("features", "NA"),
            "Survival gene-only AUC "
            + fmt(upload_metrics.get("survival", {}).get("holdout_roc_auc"))
            + "; Recurrence gene-only AUC "
            + fmt(upload_metrics.get("recurrence", {}).get("holdout_roc_auc")),
        ],
        ["Gene correlation", "METABRIC", "Gene-gene patterns", corr.get("samples", "NA"), f"{corr.get('strong_pairs_abs_ge_0_45', 'NA')} strong pairs"],
    ]
    return pd.DataFrame(rows, columns=["Part", "Dataset", "Target", "Samples", "Main result"])


def upload_metrics_table() -> pd.DataFrame:
    upload_metrics = load_json(OUT / "streamlit_upload_gene_models" / "upload_gene_models_metrics.json")
    rows = []
    for endpoint in ["survival", "recurrence"]:
        data = upload_metrics.get(endpoint, {})
        if not data:
            continue
        rows.append(
            {
                "Upload task": endpoint.title(),
                "Samples": data.get("samples", "NA"),
                "Events": data.get("events", "NA"),
                "Features": data.get("features", "NA"),
                "Selected genes": data.get("selected_genes", "NA"),
                "ROC-AUC": fmt(data.get("holdout_roc_auc")),
                "Balanced Acc": fmt(data.get("holdout_balanced_accuracy")),
                "F1": fmt(data.get("holdout_f1")),
            }
        )
    return pd.DataFrame(rows)


def build_docx() -> Path:
    DELIVERABLES.mkdir(exist_ok=True)
    doc = Document()
    style_document(doc)
    add_header_footer(doc)

    tcga = load_json(OUT / "tcga_brca_cells" / "cancer_prediction" / "cancer_holdout_metrics.json")
    bcsc = load_json(OUT / "bcsc_future_risk_prediction" / "bcsc_future_risk_metrics.json")
    subtype = load_json(OUT / "tumor_subtype_prediction" / "tumor_subtype_metrics.json")
    survival = load_json(OUT / "metabric_optimized_prognosis" / "survival_best_metrics.json")
    recurrence = load_json(OUT / "metabric_optimized_prognosis" / "recurrence_best_metrics.json")
    corr = load_json(OUT / "gene_correlation_pattern_analysis" / "gene_correlation_pattern_summary.json")
    upload_metrics = load_json(OUT / "streamlit_upload_gene_models" / "upload_gene_models_metrics.json")

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("Breast Cancer AI Project")
    r.font.name = "Calibri"
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run("Unified client report: cancer signature, future risk, subtype, prognosis, external validation, and gene patterns")
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.color.rgb = MUTED

    doc.add_heading("1. Executive Summary", level=1)
    p = doc.add_paragraph(
        "تم تنظيم المشروع في نظام واحد لتحليل سرطان الثدي باستخدام بيانات الجينات والبيانات الإكلينيكية. "
        "النظام يغطي التنبؤ بالبصمة الجينية للسرطان، تحديد نوع الورم، تقدير الخطورة المستقبلية، "
        "تحليل الـ survival والـ recurrence، والتحقق الخارجي، مع تحليل العلاقات بين الجينات."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for item in [
        "Cancer molecular signature prediction from TCGA-BRCA.",
        "Future breast cancer incidence risk percentage from BCSC screening follow-up data.",
        "Tumor subtype prediction using real METABRIC CLAUDIN subtype labels.",
        "Survival and recurrence prognosis using METABRIC real clinical labels.",
        "External validation using GEO cohorts.",
        "Gene correlation and expression pattern analysis for explainability.",
        "Streamlit demo for interactive client testing.",
        "Gene Expression Upload page where users can upload a CSV and run gene-based predictions.",
    ]:
        add_bullet(doc, item)

    add_callout(
        doc,
        "Scientific note",
        "TCGA Cancer Signature is cancer-like vs normal-like molecular detection. "
        "The future percentage risk component is handled separately using BCSC because it contains real follow-up diagnosis labels.",
    )

    doc.add_heading("2. Results Overview", level=1)
    add_table_from_df(
        doc,
        build_overview_rows(),
        widths=[1900, 1200, 2700, 1100, 2460],
        font_size=8.4,
    )

    doc.add_heading("3. TCGA Cancer Molecular Signature", level=1)
    doc.add_paragraph(
        f"The TCGA-BRCA model separates Primary Tumor samples from Solid Tissue Normal samples. "
        f"It used {tcga.get('samples', 'NA')} samples and {tcga.get('features', 'NA')} gene features. "
        f"Holdout ROC-AUC = {fmt(tcga.get('roc_auc'))}, accuracy = {fmt(tcga.get('accuracy'))}, "
        f"balanced accuracy = {fmt(tcga.get('balanced_accuracy'))}."
    )
    add_figure(doc, OUT / "tcga_brca_cells" / "cancer_prediction" / "figures" / "cancer_roc_curve.png", "Figure 1. TCGA cancer signature ROC curve")
    add_figure(doc, OUT / "tcga_brca_cells" / "cancer_prediction" / "figures" / "cancer_confusion_matrix.png", "Figure 2. TCGA cancer signature confusion matrix")
    add_figure(doc, OUT / "tcga_brca_cells" / "cancer_prediction" / "figures" / "cancer_top_genes.png", "Figure 3. Top TCGA cancer signature genes")

    doc.add_heading("4. Future Incidence Risk Percentage", level=1)
    doc.add_paragraph(
        f"The BCSC component predicts the probability of breast cancer diagnosis within one year after screening mammography. "
        f"Weighted records = {int(bcsc.get('weighted_training_records', 0) + bcsc.get('weighted_validation_records', 0)):,}. "
        f"Validation ROC-AUC = {fmt(bcsc.get('validation', {}).get('roc_auc'))}; "
        f"Brier score = {fmt(bcsc.get('validation', {}).get('brier_score'), 5)}."
    )
    add_figure(doc, OUT / "bcsc_future_risk_prediction" / "figures" / "02_bcsc_roc_curve.png", "Figure 4. BCSC future risk ROC curve")
    add_figure(doc, OUT / "bcsc_future_risk_prediction" / "figures" / "04_bcsc_calibration.png", "Figure 5. BCSC calibration")
    add_figure(doc, OUT / "bcsc_future_risk_prediction" / "figures" / "05_bcsc_top_feature_effects.png", "Figure 6. BCSC top feature effects")

    doc.add_heading("5. Tumor Subtype Prediction", level=1)
    classes = ", ".join(subtype.get("classes", []))
    doc.add_paragraph(
        f"The subtype component uses real METABRIC CLAUDIN subtype labels. "
        f"Predicted samples = {subtype.get('samples_predicted', 'NA')}; classes = {classes}. "
        f"Honest holdout accuracy = {fmt(subtype.get('holdout_accuracy_from_previous_training'))}; "
        f"holdout balanced accuracy = {fmt(subtype.get('holdout_balanced_accuracy_from_previous_training'))}."
    )
    add_figure(doc, OUT / "tumor_subtype_prediction" / "figures" / "03_subtype_confusion_matrix_all_labeled.png", "Figure 7. Tumor subtype confusion matrix")
    add_figure(doc, OUT / "tumor_subtype_prediction" / "figures" / "01_predicted_subtype_distribution.png", "Figure 8. Predicted subtype distribution")
    add_figure(doc, OUT / "tumor_subtype_prediction" / "figures" / "05_top_subtype_genes.png", "Figure 9. Top subtype genes")

    doc.add_heading("6. Prognosis: Survival and Recurrence", level=1)
    doc.add_paragraph(
        f"Survival best model: {survival.get('best_candidate', 'NA')}; "
        f"holdout ROC-AUC = {fmt(survival.get('holdout_roc_auc'))}; "
        f"balanced accuracy = {fmt(survival.get('holdout_balanced_accuracy'))}."
    )
    doc.add_paragraph(
        f"Recurrence best model: {recurrence.get('best_candidate', 'NA')}; "
        f"holdout ROC-AUC = {fmt(recurrence.get('holdout_roc_auc'))}; "
        f"balanced accuracy = {fmt(recurrence.get('holdout_balanced_accuracy'))}."
    )
    add_figure(doc, OUT / "metabric_optimized_prognosis" / "figures" / "02_optimized_performance.png", "Figure 10. Optimized prognosis performance")
    add_figure(doc, OUT / "metabric_optimized_prognosis" / "figures" / "04_holdout_roc_curves.png", "Figure 11. METABRIC holdout ROC curves")
    add_figure(doc, OUT / "metabric_optimized_prognosis" / "figures" / "06_kaplan_meier_risk_groups.png", "Figure 12. Kaplan-Meier risk groups")

    doc.add_heading("7. GEO External Validation", level=1)
    external = read_csv(OUT / "geo_external_validation" / "metabric_to_geo_external_validation_metrics.csv")
    if not external.empty:
        cols = ["label_name", "dataset", "endpoint", "samples", "events", "overlap_genes_used", "roc_auc", "balanced_accuracy", "f1"]
        ext = external[cols].copy()
        for c in ["roc_auc", "balanced_accuracy", "f1"]:
            ext[c] = ext[c].map(lambda x: fmt(x))
        add_table_from_df(
            doc,
            ext,
            widths=[1850, 900, 1600, 800, 800, 1100, 900, 1200, 800],
            font_size=7.4,
        )
    add_figure(doc, OUT / "geo_external_validation" / "figures" / "04_metabric_to_geo_external_auc.png", "Figure 13. METABRIC-to-GEO external AUC")
    add_figure(doc, OUT / "geo_external_validation" / "figures" / "05_external_roc_curves.png", "Figure 14. External validation ROC curves")

    doc.add_heading("8. Gene Correlation and Expression Patterns", level=1)
    doc.add_paragraph(
        f"METABRIC gene pattern analysis used {corr.get('samples', 'NA')} samples and "
        f"{corr.get('selected_genes', 'NA')} selected genes. "
        f"Detected strong gene-gene pairs with abs(correlation) >= 0.45: {corr.get('strong_pairs_abs_ge_0_45', 'NA')}."
    )
    pairs = read_csv(OUT / "gene_correlation_pattern_analysis" / "strong_gene_correlation_pairs_abs_ge_0_45.csv", nrows=12)
    if not pairs.empty:
        pairs = pairs[["gene_a", "gene_b", "spearman_correlation", "relationship", "gene_a_sources", "gene_b_sources"]].copy()
        pairs["spearman_correlation"] = pairs["spearman_correlation"].map(lambda x: fmt(x))
        add_table_from_df(
            doc,
            pairs,
            widths=[900, 900, 1250, 1650, 2300, 2360],
            font_size=7.3,
        )
    add_figure(doc, OUT / "gene_correlation_pattern_analysis" / "figures" / "01_gene_gene_correlation_heatmap.png", "Figure 15. Gene-gene correlation heatmap")
    add_figure(doc, OUT / "gene_correlation_pattern_analysis" / "figures" / "04_strong_gene_correlation_network.png", "Figure 16. Strong gene correlation network")
    add_figure(doc, OUT / "gene_correlation_pattern_analysis" / "figures" / "05_gene_pair_scatter_patterns.png", "Figure 17. Gene pair scatter patterns")

    doc.add_heading("9. Gene Expression Upload and Streamlit Demo", level=1)
    doc.add_paragraph(
        "The project includes patient-level CSV outputs for cancer-signature probability, prognosis percentages, subtype confidence, and GEO external event-risk probabilities. "
        "A Streamlit application was created so the client can review results interactively, enter BCSC risk-factor profiles to calculate a future risk percentage, "
        "and upload new gene-expression CSV files for gene-based testing."
    )
    add_callout(
        doc,
        "New upload feature",
        "The user can upload a CSV file containing gene-expression values. The app then attempts Cancer Detection, Tumor Subtype Prediction, gene-only Survival risk, "
        "gene-only Recurrence risk, and therapeutic/clinical support flags, depending on gene coverage in the uploaded file.",
    )
    upload_df = upload_metrics_table()
    if not upload_df.empty:
        add_table_from_df(
            doc,
            upload_df,
            widths=[1300, 850, 850, 1100, 1200, 1050, 1200, 900],
            font_size=7.8,
        )
    doc.add_paragraph(
        "Upload CSV support accepts wide format, where each row is a sample and each gene is a column, or long format, where a gene column such as Hugo_Symbol/gene is provided with one or more numeric sample/value columns. "
        "A small METABRIC-style sample file was added so the client can test the page immediately."
    )
    add_bullet(doc, "Sample upload file: sample_gene_expression_upload_metabric.csv")
    add_bullet(doc, "Cancer Detection upload requires broad TCGA-compatible gene coverage.")
    add_bullet(doc, "Subtype and gene-only prognosis upload use METABRIC-compatible gene-expression coverage.")
    add_bullet(
        doc,
        "Gene-only Survival upload AUC: "
        + fmt(upload_metrics.get("survival", {}).get("holdout_roc_auc"))
        + "; Gene-only Recurrence upload AUC: "
        + fmt(upload_metrics.get("recurrence", {}).get("holdout_roc_auc")),
    )
    add_bullet(doc, "Unified notebook: client_deliverables/breast_cancer_ai_unified_client_notebook.ipynb")
    add_bullet(doc, "Streamlit app: client_deliverables/streamlit_app.py")
    add_bullet(doc, "GitHub repository: https://github.com/Kareem-Ayman-salama/breast-cancer-ai-streamlit-demo")
    add_bullet(doc, "HTML report: client_deliverables/final_client_report.html")
    add_bullet(doc, "Markdown report: client_deliverables/final_client_report.md")

    doc.add_heading("10. Limitations and Next Improvements", level=1)
    for item in [
        "Cancer signature and future incidence risk are separate clinical questions and should be described separately.",
        "BCSC future risk is based on screening and clinical risk factors, not gene expression.",
        "The upload Survival and Recurrence models use gene expression only, so they are expected to be weaker than the optimized METABRIC combined clinical+gene models.",
        "Survival and recurrence tasks are harder and show moderate performance, which is expected for noisy prognosis labels.",
        "GEO external validation is mixed and should be reported transparently, especially weak OS transfer on GSE7390.",
        "Clinical support outputs are decision-support insights only, not direct treatment prescriptions.",
        "Next research improvements: add more external cohorts, survival-specific models, SHAP plots, and pathway enrichment.",
    ]:
        add_bullet(doc, item)

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build_docx())
